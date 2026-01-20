import os
import json
import argparse
import socket
import sqlite3
import shutil
import uuid
from concurrent.futures import ThreadPoolExecutor

from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse

from docx import Document
from openai import OpenAI


# -----------------------------
# Compatibility fix (Windows)
# Some environments may not expose socket.AF_UNIX.
# Uvicorn may reference it during startup; add a stub to prevent crash.
# -----------------------------
if not hasattr(socket, "AF_UNIX"):
    socket.AF_UNIX = 1


# ---------- dirs ----------
def data_dir() -> str:
    d = os.environ.get("MVP_DATA_DIR") or os.path.abspath("./mvp_data")
    os.makedirs(d, exist_ok=True)
    return d


def path_db() -> str:
    return os.path.join(data_dir(), "app.db")


def work_dir(task_id: str) -> str:
    d = os.path.join(data_dir(), "work", task_id)
    os.makedirs(d, exist_ok=True)
    return d


def export_dir(task_id: str) -> str:
    d = os.path.join(data_dir(), "exports", task_id)
    os.makedirs(d, exist_ok=True)
    return d


# ---------- db ----------
def db():
    conn = sqlite3.connect(path_db(), check_same_thread=False)
    conn.row_factory = sqlite3.Row
    return conn


def init_db():
    conn = db()
    cur = conn.cursor()
    cur.execute(
        """
    CREATE TABLE IF NOT EXISTS settings(
      id INTEGER PRIMARY KEY CHECK (id=1),
      base_url TEXT NOT NULL,
      api_key TEXT NOT NULL,
      model TEXT NOT NULL
    )"""
    )
    cur.execute(
        """
    CREATE TABLE IF NOT EXISTS tasks(
      id TEXT PRIMARY KEY,
      filename TEXT NOT NULL,
      source_path TEXT NOT NULL,
      work_path TEXT NOT NULL,
      direction TEXT NOT NULL,
      status TEXT NOT NULL,
      progress REAL NOT NULL,
      error TEXT
    )"""
    )
    cur.execute(
        """
    CREATE TABLE IF NOT EXISTS blocks(
      id TEXT PRIMARY KEY,
      task_id TEXT NOT NULL,
      locator TEXT NOT NULL,
      kind TEXT NOT NULL,
      order_no INTEGER NOT NULL,
      source_text TEXT NOT NULL,
      translated_text TEXT,
      status TEXT NOT NULL
    )"""
    )
    conn.commit()
    conn.close()


def get_settings():
    conn = db()
    row = conn.execute(
        "SELECT base_url, api_key, model FROM settings WHERE id=1"
    ).fetchone()
    conn.close()
    return dict(row) if row else None


def upsert_settings(base_url: str, api_key: str, model: str):
    conn = db()
    conn.execute(
        "INSERT INTO settings(id, base_url, api_key, model) VALUES(1,?,?,?) "
        "ON CONFLICT(id) DO UPDATE SET base_url=excluded.base_url, api_key=excluded.api_key, model=excluded.model",
        (base_url, api_key, model),
    )
    conn.commit()
    conn.close()


# ---------- docx extract/apply ----------
def extract_blocks(docx_path: str):
    doc = Document(docx_path)
    blocks = []
    order_no = 0

    # body paragraphs (header/footer not touched)
    for i, p in enumerate(doc.paragraphs):
        text = (p.text or "").strip()
        if text:
            blocks.append(
                {
                    "locator": f"p:{i}",
                    "kind": "paragraph",
                    "source_text": text,
                    "order_no": order_no,
                }
            )
            order_no += 1

    # tables
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                for pi, p in enumerate(cell.paragraphs):
                    text = (p.text or "").strip()
                    if text:
                        blocks.append(
                            {
                                "locator": f"t:{ti}/r:{ri}/c:{ci}/p:{pi}",
                                "kind": "table_cell_paragraph",
                                "source_text": text,
                                "order_no": order_no,
                            }
                        )
                        order_no += 1

    return blocks


def _set_paragraph_text_keep_style(paragraph, new_text: str):
    # MVP: keep paragraph object/style; write translation into first run
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for r in paragraph.runs[1:]:
            r.text = ""
    else:
        paragraph.add_run(new_text)


def apply_translations(src_docx: str, out_docx: str, locator_to_text: dict):
    doc = Document(src_docx)

    for i, p in enumerate(doc.paragraphs):
        loc = f"p:{i}"
        if loc in locator_to_text:
            _set_paragraph_text_keep_style(p, locator_to_text[loc])

    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                for pi, p in enumerate(cell.paragraphs):
                    loc = f"t:{ti}/r:{ri}/c:{ci}/p:{pi}"
                    if loc in locator_to_text:
                        _set_paragraph_text_keep_style(p, locator_to_text[loc])

    doc.save(out_docx)


def build_messages(text: str, direction: str):
    target = "English" if direction == "zh->en" else "Chinese"
    system = (
        "You are a professional translator. Output only the translation. "
        "Do not add explanations or any extra text. "
        "Preserve numbers, units, symbols, and formatting as much as possible."
    )
    user = f"Translate the following text into {target}:\n\n{text}"
    return [{"role": "system", "content": system}, {"role": "user", "content": user}]


# ---------- app ----------
init_db()
app = FastAPI(title="MVP Backend")

# Allow local UI to call local backend
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

executor = ThreadPoolExecutor(max_workers=2)


@app.get("/api/health")
def health():
    return {"ok": True}


@app.get("/api/settings")
def api_get_settings():
    s = get_settings()
    return s or {"base_url": "", "api_key": "", "model": ""}


@app.post("/api/settings")
def api_save_settings(payload: dict):
    base_url = (payload.get("base_url") or "").strip()
    api_key = (payload.get("api_key") or "").strip()
    model = (payload.get("model") or "").strip()
    if not base_url or not api_key or not model:
        raise HTTPException(400, "base_url/api_key/model required")
    upsert_settings(base_url, api_key, model)
    return {"ok": True}


@app.post("/api/tasks")
async def create_task(file: UploadFile = File(...), direction: str = Form(...)):
    if direction not in ("zh->en", "en->zh"):
        raise HTTPException(400, "direction must be zh->en or en->zh")
    if not file.filename.lower().endswith(".docx"):
        raise HTTPException(400, "only .docx supported in MVP")

    s = get_settings()
    if not s or not s["base_url"] or not s["api_key"] or not s["model"]:
        raise HTTPException(400, "please set settings first")

    task_id = f"task_{uuid.uuid4().hex}"
    wd = work_dir(task_id)
    src_path = os.path.join(wd, file.filename)

    with open(src_path, "wb") as f:
        f.write(await file.read())

    work_path = os.path.join(wd, "work.docx")
    shutil.copy2(src_path, work_path)

    blocks = extract_blocks(work_path)

    conn = db()
    conn.execute(
        "INSERT INTO tasks(id, filename, source_path, work_path, direction, status, progress, error) "
        "VALUES(?,?,?,?,?,?,?,?)",
        (task_id, file.filename, src_path, work_path, direction, "created", 0.0, None),
    )

    for b in blocks:
        conn.execute(
            "INSERT INTO blocks(id, task_id, locator, kind, order_no, source_text, translated_text, status) "
            "VALUES(?,?,?,?,?,?,?,?)",
            (
                f"blk_{uuid.uuid4().hex}",
                task_id,
                b["locator"],
                b["kind"],
                b["order_no"],
                b["source_text"],
                None,
                "pending",
            ),
        )
    conn.commit()
    conn.close()

    return {"task_id": task_id, "blocks": len(blocks)}


@app.get("/api/tasks/{task_id}")
def get_task(task_id: str):
    conn = db()
    row = conn.execute(
        "SELECT id, status, progress, error, direction FROM tasks WHERE id=?",
        (task_id,),
    ).fetchone()
    conn.close()
    if not row:
        raise HTTPException(404, "task not found")
    return dict(row)


def _translate_task(task_id: str):
    conn = db()
    try:
        task = conn.execute("SELECT * FROM tasks WHERE id=?", (task_id,)).fetchone()
        if not task:
            return

        conn.execute("UPDATE tasks SET status=?, error=? WHERE id=?", ("running", None, task_id))
        conn.commit()

        s = get_settings()
        client = OpenAI(api_key=s["api_key"], base_url=s["base_url"])
        model = s["model"]

        blocks = conn.execute(
            "SELECT * FROM blocks WHERE task_id=? ORDER BY order_no ASC", (task_id,)
        ).fetchall()

        total = len(blocks)
        done = 0

        for b in blocks:
            if b["status"] == "edited":
                done += 1
            elif b["translated_text"] and b["status"] == "translated":
                done += 1
            else:
                msgs = build_messages(b["source_text"], task["direction"])
                r = client.chat.completions.create(
                    model=model, messages=msgs, temperature=0.2
                )
                out = (r.choices[0].message.content or "").strip()
                conn.execute(
                    "UPDATE blocks SET translated_text=?, status=? WHERE id=?",
                    (out, "translated", b["id"]),
                )
                done += 1

            progress = 1.0 if total == 0 else done / total
            conn.execute("UPDATE tasks SET progress=? WHERE id=?", (progress, task_id))
            conn.commit()

        conn.execute(
            "UPDATE tasks SET status=?, progress=? WHERE id=?", ("finished", 1.0, task_id)
        )
        conn.commit()

    except Exception as e:
        conn.execute("UPDATE tasks SET status=?, error=? WHERE id=?", ("error", str(e), task_id))
        conn.commit()
    finally:
        conn.close()


@app.post("/api/tasks/{task_id}/run_translate")
def run_translate(task_id: str):
    conn = db()
    row = conn.execute("SELECT status FROM tasks WHERE id=?", (task_id,)).fetchone()
    conn.close()
    if not row:
        raise HTTPException(404, "task not found")
    if row["status"] == "running":
        return {"ok": True, "status": "running"}
    executor.submit(_translate_task, task_id)
    return {"ok": True}


@app.get("/api/tasks/{task_id}/blocks")
def list_blocks(task_id: str, offset: int = 0, limit: int = 2000):
    conn = db()
    rows = conn.execute(
        "SELECT id, order_no, status, locator, kind, source_text, translated_text "
        "FROM blocks WHERE task_id=? ORDER BY order_no ASC LIMIT ? OFFSET ?",
        (task_id, limit, offset),
    ).fetchall()
    conn.close()
    return [dict(r) for r in rows]


@app.patch("/api/tasks/{task_id}/blocks/{block_id}")
def patch_block(task_id: str, block_id: str, payload: dict):
    translated_text = payload.get("translated_text")
    if translated_text is None:
        raise HTTPException(400, "translated_text required")

    conn = db()
    row = conn.execute(
        "SELECT id FROM blocks WHERE id=? AND task_id=?", (block_id, task_id)
    ).fetchone()
    if not row:
        conn.close()
        raise HTTPException(404, "block not found")

    conn.execute(
        "UPDATE blocks SET translated_text=?, status=? WHERE id=?",
        (translated_text, "edited", block_id),
    )
    conn.commit()
    conn.close()
    return {"ok": True}


@app.get("/api/tasks/{task_id}/export")
def export_docx(task_id: str):
    conn = db()
    task = conn.execute(
        "SELECT work_path, filename FROM tasks WHERE id=?", (task_id,)
    ).fetchone()
    if not task:
        conn.close()
        raise HTTPException(404, "task not found")

    blocks = conn.execute(
        "SELECT locator, translated_text FROM blocks WHERE task_id=?", (task_id,)
    ).fetchall()
    conn.close()

    locator_to_text = {b["locator"]: b["translated_text"] for b in blocks if b["translated_text"]}

    out_path = os.path.join(export_dir(task_id), f"translated_{task['filename']}")
    apply_translations(task["work_path"], out_path, locator_to_text)

    return FileResponse(
        out_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=os.path.basename(out_path),
    )


def _atomic_write_json(path: str, data: dict):
    os.makedirs(os.path.dirname(path), exist_ok=True)
    tmp = path + ".tmp"
    with open(tmp, "w", encoding="utf-8") as f:
        json.dump(data, f)
        f.flush()
        os.fsync(f.fileno())
    os.replace(tmp, path)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--host", default="127.0.0.1")
    parser.add_argument("--port-file", required=True)
    args = parser.parse_args()

    # Pick a free TCP port (do NOT use uvicorn fd-mode; improves Windows compatibility)
    s = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
    s.bind((args.host, 0))
    port = s.getsockname()[1]
    s.close()

    _atomic_write_json(args.port_file, {"host": args.host, "port": port, "pid": os.getpid()})

    import uvicorn
    uvicorn.run(
        app,
        host=args.host,
        port=port,
        log_level="info",
        access_log=False,
        use_colors=False,
    )


if __name__ == "__main__":
    main()
